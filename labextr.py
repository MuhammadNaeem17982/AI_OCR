import camelot
import re
import os
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import spacy
from docx import Document
import openai

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
nlp = spacy.load("en_core_web_trf")
openai.api_key = ""

def generate_explanation_lab(text):
    prompt = "I have a blood test with the following results:\n" + text + "\nAnalyze the results and provide detailed explanation."
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a medical expert analyzing blood test results."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000
    )
    return response["choices"][0]["message"]["content"]

def extract_tables_from_pdf_lab(pdf_path):
    tables = camelot.read_pdf(pdf_path, pages='all', flavor='stream')
    return tables

def clean_table_text_lab(table_text):
    cleaned_lines = []
    for line in table_text.split("\n"):
        if not re.search(r'Printed by|Page \d+ of \d+|Confidential Report|End of report|IMPORTANT|Sample Received|Investigation/Test', line) and line.strip():
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def preprocess_image_lab(image_path):
    img = Image.open(image_path)
    img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
    img = img.convert('L')
    img = img.filter(ImageFilter.MedianFilter())
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(2)
    enhancer = ImageEnhance.Brightness(img)
    img = enhancer.enhance(1.5)
    img = img.filter(ImageFilter.SHARPEN)
    return img

def extract_text_from_image_lab(image_path):
    img = preprocess_image_lab(image_path)
    custom_config = r'--oem 3 --psm 4'
    text = pytesseract.image_to_string(img, config=custom_config)
    return text.strip()

def remove_personal_info_lab(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "ORG", "LOC", "FAC", "NORP", "DATE"]:
            text = text.replace(ent.text, "[REDACTED]")
    text = re.sub(r'\b(Male|Female|Sex|Gender)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(Age|Years|Yrs)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(Dr\.|Doctor|MD|Pathologist)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'http\S+|www\S+', '[REDACTED]', text)
    text = re.sub(r'\b\d{1,2}:\d{2}\s?(AM|PM|am|pm)?\b', '[REDACTED]', text)
    return text

def extract_text_from_docx_lab(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_text))
    return '\n'.join(full_text)

def process_file_lab(file_path):
    _, file_extension = os.path.splitext(file_path)
    result = ""
    if file_extension.lower() == '.pdf':
        tables = extract_tables_from_pdf_lab(file_path)
        for i, table in enumerate(tables):
            cleaned_table = clean_table_text_lab(table.df.to_string(index=False))
            result += f"\n=== Table {i+1} ===\n{cleaned_table}"
    elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
        text = extract_text_from_image_lab(file_path)
        result += text
    elif file_extension.lower() == '.docx':
        result = extract_text_from_docx_lab(file_path)
    else:
        result = "Unsupported file format"
    return result

if __name__ == "__main__":
    file_path = r"C:\Users\38066\Dropbox\Team AI\Roman\test\lab - dradem.docx"
    extracted_text = process_file_lab(file_path)
    text_no_personal_info = remove_personal_info_lab(extracted_text)
    explanation = generate_explanation_lab(text_no_personal_info)
    print(text_no_personal_info)
    print(explanation)