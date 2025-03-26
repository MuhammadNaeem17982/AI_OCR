import re
import os
from pdfminer.high_level import extract_text
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
import spacy
import openai

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
nlp = spacy.load("en_core_web_md")
openai.api_key = ''

def clean_text(text):
    cleaned_lines = []
    for line in text.split("\n"):
        if not re.search(r'Printed by|Page \d+ of \d+|Confidential Report|End of report|IMPORTANT|Sample Received|Investigation/Test', line) and line.strip():
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def remove_personal_info(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "ORG", "LOC", "FAC", "NORP", "DATE", "TIME", "LANGUAGE"]:
            text = text.replace(ent.text, "[REDACTED]")
    text = re.sub(r'\b\d{10}\b', '[REDACTED]', text)
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', '[REDACTED]', text)
    text = re.sub(r'\b(?:Tel|Fax|Tel\\Fax)?[:\s]*\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(?:\+?\d{1,3})?[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}\b', '[REDACTED]', text)
    text = re.sub(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '[REDACTED]', text)
    text = re.sub(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b\d{4}-\d{2}-\d{2}\b', '[REDACTED]', text)
    text = re.sub(r'\bExamination date[-:]\s*\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(?:Repromed\+? Hospital|MEDIGROUP Slavija|Diagnostic Center|Medical Imaging|KOOROSH DIAGNOSTIC CENTER)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(Male|Female|Sex|Gender)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(Age|Years|Yrs|yo)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(Dr\.|Doctor|MD|Pathologist|Radiologist|E\.B\.C\.R)\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bΑΜΚΑ\s*\d{11}\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'http\S+|www\S+', '[REDACTED]', text)
    text = re.sub(r'\b\d{1,2}:\d{2}\s?(AM|PM|am|pm)?\b', '[REDACTED]', text)
    text = re.sub(r'\bCompany number:\s*\d+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bTIN:\s*\d+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bBA:\s*\d+[-\d]+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bProtocol number:\s*[A-Z0-9\-\/]+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bOrder[:\s#]*\d+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\bLinked Orders\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b[Kk]neginje\s+[A-Za-z\s]+\s+broj\s+\d+\b', '[REDACTED]', text)
    text = re.sub(r'\bSworn-In\s+[A-Za-z\s\-]+\b', '[REDACTED]', text, flags=re.IGNORECASE)
    text = re.sub(r'\b\w+:\s*\[REDACTED\]', '[REDACTED]', text)
    return text

def preprocess_image(image_path):
    img = Image.open(image_path)
    img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
    img = img.convert('L')
    img = img.filter(ImageFilter.MedianFilter())
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(2)
    enhancer = ImageEnhance.Brightness(img)
    img = enhancer.enhance(1.5)
    img = img.filter(ImageFilter.SHARPEN)
    return img

def extract_text_from_image(image_path):
    img = preprocess_image(image_path)
    custom_config = r'--oem 3 --psm 4'
    text = pytesseract.image_to_string(img, config=custom_config)
    return text.strip()

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_text))
    return '\n'.join(full_text)

def process_file(file_path):
    _, file_extension = os.path.splitext(file_path)
    result = ""
    if file_extension.lower() == '.pdf':
        result = extract_text(file_path)
    elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
        result = extract_text_from_image(file_path)
    elif file_extension.lower() == '.docx':
        result = extract_text_from_docx(file_path)
    else:
        result = "Unsupported file format"
    return result

def generate_explanation(text):
    prompt = "I have this document:\n" + text + "\n Analyse it and give me medcial findings."
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a medical expert ."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000
    )
    return response["choices"][0]["message"]["content"]

if __name__ == "__main__":
    file_path = r"C:\Users\38066\Dropbox\Team AI\Roman\test\ct - cf - cens.pdf"
    extracted_text = process_file(file_path)
    cleaned_text = clean_text(extracted_text)
    text_no_personal_info = remove_personal_info(cleaned_text)
    explanation = generate_explanation(text_no_personal_info)
    print(explanation)
    print(text_no_personal_info)